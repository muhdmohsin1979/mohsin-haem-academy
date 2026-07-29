#!/usr/bin/env python3
"""Validate the exact MCL v2.0 production candidate."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import fitz
from bs4 import BeautifulSoup
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PREVIEW = ROOT / "docs" / "mcl-v2" / "preview"
MCL = ROOT / "guidelines" / "mcl"
MANIFEST = MCL / "release-manifest-v2.0.json"
REVIEWED_COMMIT = "763863cf3b89d207fce3bd29c4df076e0386ab92"
REVIEWED_TREE = "e4987ea5d1ee89bf0d9b307062299eb106d26af0"
REVIEWED_MANIFEST_SHA256 = "be160f955203e33b3a72e4b9829328358568bb8ba5337eb9922fb2f8bdff95fb"
ARTEFACTS = [
    MCL / "index.html",
    MCL / "guideline-v2.0.docx",
    MCL / "guideline-v2.0.pdf",
    MCL / "quickref-v2.0.docx",
    MCL / "quickref-v2.0.pdf",
    MCL / "algorithm-v2.0.svg",
    MCL / "algorithm-v2.0.excalidraw",
]
FORBIDDEN = (
    "WORKING PREVIEW",
    "NOT FOR CLINICAL USE",
    "Clinical and pharmacy review pending",
    "no publication authority",
    "non-public working",
)


def strict_json(path: Path):
    def unique(pairs):
        result = {}
        for key, value in pairs:
            if key in result:
                raise AssertionError(f"Duplicate JSON key {key!r}: {path}")
            result[key] = value
        return result
    return json.loads(path.read_text(encoding="utf-8"), object_pairs_hook=unique)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def docx_text(path: Path) -> str:
    document = Document(str(path))
    chunks = [p.text for p in document.paragraphs]
    for section in document.sections:
        for container in (section.header, section.footer):
            chunks.extend(p.text for p in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        chunks.extend(p.text for p in cell.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    with zipfile.ZipFile(path) as archive:
        for name in ("docProps/core.xml", "docProps/app.xml", "docProps/custom.xml"):
            if name in archive.namelist():
                chunks.append(archive.read(name).decode("utf-8", errors="replace"))
    return "\n".join(chunks)


def readable(path: Path) -> tuple[str, int | None]:
    if path.suffix == ".docx":
        return docx_text(path), None
    if path.suffix == ".pdf":
        with fitz.open(path) as document:
            pages = [str(page.get_text()) for page in document]
            if any(len(page.strip()) < 80 for page in pages):
                raise AssertionError(f"Blank or near-blank page: {path.name}")
            return "\n".join(pages), document.page_count
    return path.read_text(encoding="utf-8"), None


def svg_semantics(path: Path) -> dict[str, str]:
    root = ET.parse(path).getroot()
    values = {}
    for element in root.iter():
        semantic_id = element.attrib.get("data-semantic-id")
        if semantic_id:
            values[semantic_id] = re.sub(r"\s+", " ", " ".join(element.itertext())).strip()
    return values


def excalidraw_semantics(path: Path) -> dict[str, str]:
    scene = strict_json(path)
    return {
        element["customData"]["semanticId"]: re.sub(r"\s+", " ", element.get("text", "")).strip()
        for element in scene["elements"]
        if isinstance(element, dict)
        and isinstance(element.get("customData"), dict)
        and element["customData"].get("semanticId")
    }


def normalise_release_control(text: str) -> str:
    return text.replace(
        "Clinical and pharmacy review pending — no publication authority.",
        "[RELEASE-REVIEW-STATE]",
    ).replace(
        "Independent clinical review and release pharmacy verification complete.",
        "[RELEASE-REVIEW-STATE]",
    )


def validate() -> None:
    manifest = strict_json(MANIFEST)
    if manifest.get("status") != "AUTHORISED FOR PUBLICATION 29 JULY 2026":
        raise AssertionError("Production manifest status is not authorised")
    if manifest.get("publication_authority") is not True:
        raise AssertionError("Production manifest lacks publication authority")
    reviewed = manifest.get("reviewed_preview")
    expected_reviewed = {
        "commit": REVIEWED_COMMIT,
        "tree": REVIEWED_TREE,
        "manifest_sha256": REVIEWED_MANIFEST_SHA256,
    }
    if reviewed != expected_reviewed:
        raise AssertionError("Production candidate is not bound to the reviewed preview")
    if manifest.get("clinical_change_from_reviewed_preview") != "NONE; release-control presentation only":
        raise AssertionError("Clinical-change boundary is not fail-closed")
    if sha256(PREVIEW / "build-manifest-working.json") != REVIEWED_MANIFEST_SHA256:
        raise AssertionError("Reviewed preview manifest changed")

    records = manifest.get("artefacts")
    if not isinstance(records, dict) or set(records) != {path.name for path in ARTEFACTS}:
        raise AssertionError("Production artefact family is incomplete")
    texts = {}
    for path in ARTEFACTS:
        text, pages = readable(path)
        texts[path.name] = text
        record = records[path.name]
        if record.get("sha256") != sha256(path) or record.get("bytes") != path.stat().st_size:
            raise AssertionError(f"Manifest mismatch: {path.name}")
        if pages is not None and record.get("pages") != pages:
            raise AssertionError(f"Manifest page mismatch: {path.name}")
        for phrase in FORBIDDEN:
            if phrase.casefold() in text.casefold():
                raise AssertionError(f"Stale preview state in {path.name}: {phrase}")

    if records["guideline-v2.0.pdf"].get("pages") != 19:
        raise AssertionError("Guideline pagination changed from reviewed preview")
    if records["quickref-v2.0.pdf"].get("pages") != 3:
        raise AssertionError("Quick-reference pagination changed from reviewed preview")
    for pdf_name, docx_name in (
        ("guideline-v2.0.pdf", "guideline-v2.0.docx"),
        ("quickref-v2.0.pdf", "quickref-v2.0.docx"),
    ):
        if records[pdf_name].get("source_docx_sha256") != records[docx_name].get("sha256"):
            raise AssertionError(f"Frozen PDF is not bound to its canonical DOCX: {pdf_name}")

    preview_html = BeautifulSoup((PREVIEW / "index.html").read_text(encoding="utf-8"), "html.parser")
    production_html = BeautifulSoup((MCL / "index.html").read_text(encoding="utf-8"), "html.parser")
    preview_units = [str(node) for node in preview_html.select("section[data-clinical-unit]")]
    production_units = [str(node) for node in production_html.select("section[data-clinical-unit]")]
    if preview_units != production_units or len(production_units) != 12:
        raise AssertionError("Production HTML changed reviewed clinical units")

    preview_svg = svg_semantics(PREVIEW / "algorithm-working.svg")
    production_svg = svg_semantics(MCL / "algorithm-v2.0.svg")
    preview_excal = excalidraw_semantics(PREVIEW / "algorithm-working.excalidraw")
    production_excal = excalidraw_semantics(MCL / "algorithm-v2.0.excalidraw")
    mutable_ids = {"header.preview", "footer.warning"}
    for semantic_id in set(preview_svg) - mutable_ids:
        if normalise_release_control(production_svg.get(semantic_id, "")) != normalise_release_control(preview_svg[semantic_id]):
            raise AssertionError(f"Production SVG clinical semantic changed: {semantic_id}")
    for semantic_id in set(preview_excal) - mutable_ids:
        if normalise_release_control(production_excal.get(semantic_id, "")) != normalise_release_control(preview_excal[semantic_id]):
            raise AssertionError(f"Production Excalidraw clinical semantic changed: {semantic_id}")
    if production_svg != production_excal:
        raise AssertionError("Production SVG and Excalidraw semantic models differ")

    required = ("MHA-MCL-2026-v2.0",)
    for name, text in texts.items():
        for phrase in required:
            if phrase not in text:
                raise AssertionError(f"Missing controlled identity in {name}: {phrase}")
        if "28 July 2026" not in text and "2026-07-28" not in text:
            raise AssertionError(f"Missing controlled evidence/access cut-off in {name}")
    print("MCL v2.0 production-candidate validation: PASS artefacts=7 clinical_units=12 guideline_pages=19 quickref_pages=3")


if __name__ == "__main__":
    validate()
