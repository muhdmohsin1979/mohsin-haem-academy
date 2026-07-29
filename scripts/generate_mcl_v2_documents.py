#!/usr/bin/env python3
"""Generate non-public MCL v2.0 DOCX review artefacts.

The generator is deliberately PREVIEW-only. It consumes the controlled clinical
source, access matrix and evidence ledger, and cannot add publication authority.
PDF conversion is handled separately by LibreOffice.
"""

from __future__ import annotations

import json
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "sources" / "mcl" / "source-v2.0.html"
STATUS = ROOT / "sources" / "mcl" / "status-matrix-v2.0.json"
STATE = ROOT / "sources" / "mcl" / "release-state-v2.0.json"
EVIDENCE = ROOT / "docs" / "mcl-v2" / "evidence-ledger.json"
DEFAULT_OUTPUT = ROOT / "docs" / "mcl-v2" / "preview"

NAVY = "1B2A4A"
RED = "C41E3A"
LIGHT_BLUE = "EAF0F7"
LIGHT_RED = "FBECEF"
LIGHT_GREY = "F3F4F6"
MID_GREY = "6B7280"
BORDER = "D7DCE2"
WHITE = "FFFFFF"
DOC_CODE = "MHA-MCL-2026-v2.0"
CUT_OFF = "28 July 2026"
PREVIEW_WARNING = "CONTROLLED WORKING PREVIEW — NOT FOR CLINICAL USE"
CONTROLLED_ZIP_TIME = (2026, 7, 28, 7, 0, 0)
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def load_json(path: Path) -> dict[str, object]:
    def reject_duplicate(pairs: list[tuple[str, object]]) -> dict[str, object]:
        result: dict[str, object] = {}
        for key, value in pairs:
            if key in result:
                raise ValueError(f"Duplicate JSON key in {path}: {key}")
            result[key] = value
        return result

    value = json.loads(path.read_text(encoding="utf-8"), object_pairs_hook=reject_duplicate)
    if not isinstance(value, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return value


def canonicalise_docx(path: Path) -> None:
    """Remove template residue and rewrite the OOXML container deterministically."""
    with zipfile.ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}

    for name in list(members):
        if name.startswith("customXml/") or name == "docProps/thumbnail.jpeg":
            del members[name]

    for relationship_path in ("_rels/.rels", "word/_rels/document.xml.rels"):
        if relationship_path not in members:
            continue
        root = ET.fromstring(members[relationship_path])
        for relationship in list(root):
            target = relationship.attrib.get("Target", "").casefold()
            relation_type = relationship.attrib.get("Type", "").casefold()
            if "customxml" in target or "customxml" in relation_type or "thumbnail" in target or "thumbnail" in relation_type:
                root.remove(relationship)
        ET.register_namespace("", REL_NS)
        members[relationship_path] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    content_types = "[Content_Types].xml"
    root = ET.fromstring(members[content_types])
    for item in list(root):
        part_name = item.attrib.get("PartName", "")
        if part_name.startswith("/customXml/") or part_name == "/docProps/thumbnail.jpeg":
            root.remove(item)
    ET.register_namespace("", CONTENT_TYPE_NS)
    members[content_types] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    members["docProps/app.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        '<Application>Mohsin Haematology Academy controlled generator</Application>'
        '<AppVersion>2.0</AppVersion><DocSecurity>0</DocSecurity><ScaleCrop>false</ScaleCrop>'
        '<Company>Mohsin Haematology Academy</Company><LinksUpToDate>false</LinksUpToDate>'
        '<SharedDoc>false</SharedDoc><HyperlinksChanged>false</HyperlinksChanged>'
        '</Properties>'
    ).encode("utf-8")

    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as output:
        for name in sorted(members):
            info = zipfile.ZipInfo(name, CONTROLLED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o100644 << 16
            output.writestr(info, members[name])
    temporary.replace(path)


def require_preview_state(state: dict[str, object]) -> None:
    if state.get("state") != "PREVIEW":
        raise ValueError("MCL v2.0 document generation supports PREVIEW only")
    if state.get("publication_authority") is not False:
        raise ValueError("Preview document generation requires publication_authority=false")
    for gate in ("independent_clinical_review", "pharmacy_verification"):
        if state.get(gate) != "PENDING":
            raise ValueError(f"Preview document generation requires {gate}=PENDING")


def plain_markdown(text: str) -> str:
    text = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("*", "").replace("`", "")
    return re.sub(r"\s+", " ", text).strip()


def node_text(node: Tag) -> str:
    parts: list[str] = []

    def walk(item: Tag | NavigableString) -> None:
        if isinstance(item, NavigableString):
            parts.append(str(item))
            return
        if item.name == "br":
            parts.append("\n")
            return
        for child in item.children:
            walk(child)
        if item.name in {"p", "li", "div", "td", "th"}:
            parts.append(" ")

    walk(node)
    text = "".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, colour: str = BORDER, size: int = 6) -> None:
    props = cell._tc.get_or_add_tcPr()
    borders = props.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), colour)


def keep_with_next(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def keep_row_together(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document, subtitle: str, compact: bool = False) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    margin = 1.35 if compact else 1.7
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.3 if compact else 9.2)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(3.2 if compact else 5)
    normal.paragraph_format.line_spacing = 1.05 if compact else 1.08

    for name, size, colour in (
        ("Title", 23 if compact else 25, NAVY),
        ("Heading 1", 16 if compact else 17, NAVY),
        ("Heading 2", 13 if compact else 13, RED),
        ("Heading 3", 10.5 if compact else 10.5, NAVY),
    ):
        style = document.styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(6 if compact else 10)
        style.paragraph_format.space_after = Pt(3 if compact else 5)

    if "MHA Callout" not in [style.name for style in document.styles]:
        style = document.styles.add_style("MHA Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Georgia"
        style.font.size = Pt(8.5 if compact else 8.8)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_after = Pt(3)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.8))
    table.cell(0, 0).text = "MOHSIN HAEMATOLOGY ACADEMY"
    table.cell(0, 1).text = "WORKING PREVIEW v2.0"
    for index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, NAVY if index == 0 else RED)
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f"{DOC_CODE} | {subtitle} | cut-off {CUT_OFF} | non-public working preview"
    paragraph.runs[0].font.size = Pt(8)
    paragraph.runs[0].font.color.rgb = RGBColor.from_string(MID_GREY)
    add_page_number(footer.add_paragraph())

    properties = document.core_properties
    properties.title = f"MCL v2.0 — {subtitle} — working preview"
    properties.subject = "Non-public specialist clinical working draft"
    properties.author = "Dr Muhammad Mohsin, Consultant Haematologist"
    properties.keywords = "mantle cell lymphoma, MCL, working preview"
    properties.comments = "Clinical, pharmacy and publication review pending. No publication authority."
    controlled_time = datetime(2026, 7, 28, 7, 0, tzinfo=timezone.utc)
    properties.created = controlled_time
    properties.modified = controlled_time


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(RED)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{DOC_CODE}\nEvidence and access cut-off: {CUT_OFF}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_RED)
    set_cell_border(cell, RED, 10)
    warning = cell.paragraphs[0]
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run(PREVIEW_WARNING)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(RED)
    paragraph = cell.add_paragraph("Clinical and pharmacy review pending. No publication authority. Do not use for treatment, prescribing, consent, referral or commissioning decisions.")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = document.styles["MHA Callout"]


def render_source_section(document: Document, section: Tag) -> None:
    for child in section.children:
        if not isinstance(child, Tag):
            continue
        if child.name == "h2":
            paragraph = document.add_paragraph(node_text(child), style="Heading 1")
            keep_with_next(paragraph)
        elif child.name == "h3":
            paragraph = document.add_paragraph(node_text(child), style="Heading 2")
            keep_with_next(paragraph)
        elif child.name == "p":
            document.add_paragraph(node_text(child))
        elif child.name in {"ul", "ol"}:
            style = "List Bullet" if child.name == "ul" else "List Number"
            for item in child.find_all("li", recursive=False):
                paragraph = document.add_paragraph(node_text(item), style=style)
                paragraph.paragraph_format.space_after = Pt(2)


def add_access_entry(document: Document, treatment: dict[str, object]) -> None:
    heading = document.add_paragraph(str(treatment["title"]), style="Heading 2")
    keep_with_next(heading)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = (
        ("Population", "population"),
        ("Clinical evidence", "evidence_position"),
        ("Marketing authorisation", "marketing_authorisation"),
        ("NICE / HTA", "nice_status"),
        ("NHS England access", "england_access"),
        ("Devolved-nation notes", "devolved_notes"),
        ("Regimen", "regimen"),
        ("Dose, schedule and duration", "dose_schedule_duration"),
        ("Administration/monitoring boundary", "administration_monitoring_boundary"),
        ("Pharmacy evidence status", "pharmacy_evidence_status"),
        ("Provisional public wording", "public_wording"),
    )
    for label, key in rows:
        row = table.add_row()
        keep_row_together(row)
        row.cells[0].text = label
        row.cells[1].text = str(treatment[key])
        set_cell_shading(row.cells[0], LIGHT_GREY)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                keep_with_next(paragraph)
    urls = treatment["official_source_urls"]
    if isinstance(urls, list):
        paragraph = document.add_paragraph("Official sources: " + " · ".join(str(url) for url in urls))
        paragraph.style = document.styles["MHA Callout"]


def build_guideline(output: Path, source_html: str, treatments: list[dict[str, object]], evidence_records: list[dict[str, object]]) -> None:
    document = Document()
    configure_document(document, "Full specialist working draft")
    add_title(document, "Mantle Cell Lymphoma", "Unified clinical guideline — v2.0 working draft")
    document.add_heading("Control status", level=1)
    document.add_paragraph("This build is generated from controlled source material in PREVIEW state. Clinical review, pharmacy verification, publication review, owner approval, preview verification and publication authorisation remain pending.")

    soup = BeautifulSoup(source_html, "html.parser")
    for section in soup.find_all("section", attrs={"data-clinical-unit": True}, recursive=False):
        render_source_section(document, section)

    heading = document.add_paragraph("13. Working regulatory and access matrix", style="Heading 1")
    keep_with_next(heading)
    document.add_paragraph("Marketing authorisation, HTA recommendation, NHS commissioning and trial or exceptional access are separate determinations. All entries are time-sensitive and require live recheck before candidate freeze.")
    for treatment in treatments:
        add_access_entry(document, treatment)

    document.add_heading("Evidence references", level=1)
    document.add_paragraph("Targeted independently verified evidence update; not a formal systematic review. Scientific extraction is abstract-only except where the controlled evidence ledger states otherwise.")
    for record in evidence_records:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(f"{record['id']}: ").bold = True
        paragraph.add_run(plain_markdown(str(record["bibliographic_identity_markdown"])))
        paragraph.add_run(f" | Design/population: {record['design_population']} | Verified extraction: {record['abstract_supported_result']} | Integrity: {record['verification_integrity']}")

    document.add_heading("Release-control record", level=1)
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Document code", DOC_CODE),
        ("State", "PREVIEW"),
        ("Evidence/access cut-off", CUT_OFF),
        ("Clinical review", "PENDING"),
        ("Pharmacy review", "PENDING"),
        ("Publication review", "PENDING"),
        ("Publication authority", "FALSE"),
    ):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GREY)
        for cell in row.cells:
            set_cell_border(cell)
    document.save(output)
    canonicalise_docx(output)


def add_compact_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, NAVY)
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(9.4)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = value
            set_cell_border(cell)
            if row_index % 2:
                set_cell_shading(cell, LIGHT_GREY)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.2)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1.5)


def part_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="Heading 1")
    keep_with_next(paragraph)


def build_quickref(output: Path, treatments: list[dict[str, object]]) -> None:
    by_id = {str(item["id"]): item for item in treatments}
    document = Document()
    configure_document(document, "Three-part quick reference", compact=True)
    add_title(document, "MCL quick reference", "v2.0 specialist working draft")

    part_heading(document, "PART 1 OF 3 — FIRST-LINE ORIENTATION")
    document.add_heading("Before treatment", level=2)
    add_bullets(document, [
        "Confirm integrated diagnosis and record blastoid or pleomorphic morphology, Ki-67 and TP53 status where technically feasible.",
        "Define symptoms, tempo, organ compromise, cytopenias, performance status, comorbidity, frailty and treatment goals.",
        "Separate clinical evidence from licence, NICE/HTA status and operational access.",
    ])
    add_compact_table(document, ["Clinical state", "Working orientation", "Access boundary"], [
        ["Younger / treatment-fit", "TRIANGLE changes the evidence architecture: the studied ibrutinib-containing non-ASCT strategy had mature efficacy; adding ASCT did not improve FFS.", "GB-authorised regimen; NICE recommendation remains draft and negative. No final national entitlement at cut-off."],
        ["Older / transplant-ineligible", "Choose through disease risk, frailty, renal, infection, cardiac, bleeding and interaction assessment. ECHO and ENRICH inform evidence but do not erase access constraints.", "VR-CAP has TA370 in its exact population. Other modern regimens require separate live access confirmation."],
        ["TP53-mutated / high risk", "Early specialist and trial discussion. Single-arm novel combinations show activity but do not prove comparative superiority or elimination of TP53 risk.", "Do not present trial or early-access routes as routine commissioned treatment."],
        ["Observation", "Selected asymptomatic low-volume, indolent or non-nodal disease may be observed after diagnostic confidence and MDT review.", "Record review interval and objective reassessment triggers."],
    ])
    document.add_paragraph("Clinical and pharmacy review pending. This quick reference does not contain prescribing schedules and cannot replace the SmPC or local SACT protocol.", style="MHA Callout")

    document.add_page_break()
    part_heading(document, "PART 2 OF 3 — RELAPSED OR REFRACTORY PATHWAY")
    document.add_heading("At every relapse", level=2)
    add_bullets(document, [
        "Record every prior regimen and class, response depth and duration, treatment-free interval and reason for stopping.",
        "Distinguish covalent-BTKi intolerance from progression.",
        "Reassess tempo, tissue diagnosis where needed, performance status, organ function, infection history and trial eligibility.",
        "Refer early for cellular-therapy assessment where appropriate; account for manufacturing time and bridging risk.",
    ])
    add_compact_table(document, ["State", "Working pathway", "Do not"], [
        ["Exactly one previous line", "Ibrutinib TA502 or zanubrutinib TA1081 in England when exact live criteria are met.", "Do not generalise the NICE one-line restriction to the broader licences."],
        ["Covalent-BTKi intolerance", "Consider another clinically appropriate route according to retained sensitivity, toxicity, comorbidity and access.", "Do not call intolerance resistant disease."],
        ["Covalent-BTKi progression", "Assess CAR-T eligibility and trial options early. Consider non-covalent BTKi evidence only with a confirmed access route.", "Do not infer routine pirtobrutinib funding from its MCL licence or from CLL access."],
        ["Post-BTKi, cellular therapy eligible", "Brexu-cel remains TA677 CDF managed access when live national and panel criteria are met.", "Do not describe TA677 as unrestricted baseline commissioning."],
        ["No routine national route", "Use trial, early-access or exceptional pathways only when explicitly confirmed and governed.", "Do not equate absence of a national route with absolute unavailability."],
    ])
    document.add_paragraph("Do not use unadjusted cross-trial response rates to rank pirtobrutinib, cellular therapy, bispecific antibodies or investigational BCL2 inhibition.", style="MHA Callout")

    document.add_page_break()
    part_heading(document, "PART 3 OF 3 — ACCESS AND SAFETY CONTROLS")
    routine_rows: list[list[str]] = []
    for treatment_id in ("vr-cap-first-line", "ibrutinib-rr-one-line", "zanubrutinib-rr-one-line", "brexu-cel-managed-access"):
        item = by_id[treatment_id]
        routine_rows.append([str(item["title"]), str(item["population"]), str(item["public_wording"])])
    add_compact_table(document, ["Treatment", "Exact population", "England position at cut-off"], routine_rows)
    document.add_heading("Time-sensitive non-routine or unfinished routes", level=2)
    add_bullets(document, [
        "TRIANGLE ibrutinib: licensed; current NICE recommendation is draft and negative, not final guidance.",
        "Acalabrutinib–BR: licensed for previously untreated adults not eligible for ASCT, but no demonstrated final positive national MCL route at cut-off.",
        "Acalabrutinib monotherapy: licensed for relapsed or refractory MCL not previously treated with a BTKi, but no demonstrated national England MCL commissioning route.",
        "Pirtobrutinib: licensed after previous BTKi treatment, but no demonstrated national England MCL commissioning route at cut-off.",
        "Liso-cel: licensed after at least two systemic lines including a BTKi, but no demonstrated national England MCL commissioning route at cut-off.",
        "Ibrutinib–venetoclax: no current MCL marketing authorisation; NICE appraisal suspended.",
        "Glofitamab and epcoritamab: do not extrapolate licences or funding from other lymphoma indications.",
    ])
    document.add_heading("Pharmacy and supportive-care hold points", level=2)
    add_bullets(document, [
        "Use the current SmPC and local SACT protocol for formulation, dose, schedule, modifications and administration.",
        "Complete infection, vaccination, viral-screening, interaction, bleeding, cardiac, renal and tumour-lysis assessment.",
        "Use dedicated CAR-T and bispecific toxicity-management pathways.",
        "Human pharmacy verification is mandatory before any publication candidate is frozen.",
    ])
    document.add_paragraph("Clinical and pharmacy review pending — no publication authority. Recheck official national sources before treatment and before any release decision.", style="MHA Callout")
    document.save(output)
    canonicalise_docx(output)


def build_documents(output_dir: Path = DEFAULT_OUTPUT, state_path: Path = STATE) -> dict[str, Path]:
    source_html = SOURCE.read_text(encoding="utf-8")
    if "[VERIFY]" in source_html or re.search(r"<script\b|\son[a-z]+\s*=", source_html, re.IGNORECASE):
        raise ValueError("Canonical MCL source contains unresolved or active markup")
    state = load_json(state_path)
    require_preview_state(state)
    matrix = load_json(STATUS)
    ledger = load_json(EVIDENCE)
    treatments_value = matrix.get("treatments")
    records_value = ledger.get("records")
    if not isinstance(treatments_value, list) or len(treatments_value) < 10:
        raise ValueError("Treatment-access matrix is incomplete")
    if not isinstance(records_value, list) or len(records_value) != ledger.get("record_count"):
        raise ValueError("Evidence record count is inconsistent")
    treatments = [item for item in treatments_value if isinstance(item, dict)]
    records = [item for item in records_value if isinstance(item, dict)]
    if len(treatments) != len(treatments_value) or len(records) != ledger.get("record_count"):
        raise ValueError("Malformed treatment or evidence records")

    output_dir.mkdir(parents=True, exist_ok=True)
    guideline = output_dir / "guideline-working.docx"
    quickref = output_dir / "quickref-working.docx"
    build_guideline(guideline, source_html, treatments, records)
    build_quickref(quickref, treatments)
    return {"guideline": guideline, "quickref": quickref}


def main() -> int:
    outputs = build_documents()
    for path in outputs.values():
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
