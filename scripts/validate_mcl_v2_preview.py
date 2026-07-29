#!/usr/bin/env python3
"""Validate the non-public MCL v2.0 multi-format working preview."""

from __future__ import annotations

import hashlib
import json
import sys
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path

import fitz
from bs4 import BeautifulSoup
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.build_mcl_v2_preview import CONTROLLED_INPUTS, build_multiformat_preview
from scripts.generate_mcl_v2_algorithm import canonical_semantic_texts

PREVIEW = ROOT / "docs" / "mcl-v2" / "preview"
SOURCE = ROOT / "sources" / "mcl" / "source-v2.0.html"
STATUS = ROOT / "sources" / "mcl" / "status-matrix-v2.0.json"
STATE = ROOT / "sources" / "mcl" / "release-state-v2.0.json"
MANIFEST = PREVIEW / "build-manifest-working.json"


def strict_json_load(path: Path) -> object:
    def reject_duplicates(pairs: list[tuple[str, object]]) -> dict[str, object]:
        result: dict[str, object] = {}
        for key, value in pairs:
            if key in result:
                raise ValueError(f"Duplicate JSON key {key!r} in {path}")
            result[key] = value
        return result

    return json.loads(path.read_text(encoding="utf-8"), object_pairs_hook=reject_duplicates)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def docx_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def pdf_text_and_pages(path: Path) -> tuple[str, int]:
    with fitz.open(path) as document:
        chunks: list[str] = []
        for number, page in enumerate(document, start=1):
            text = page.get_text()
            if not text.strip():
                raise AssertionError(f"{path.name} page {number} is blank")
            for block in page.get_text("blocks"):
                x0, y0, x1, y1 = block[:4]
                if x0 < -1 or y0 < -1 or x1 > page.rect.width + 1 or y1 > page.rect.height + 1:
                    raise AssertionError(f"{path.name} page {number} contains text outside page bounds")
            chunks.append(text)
        return "\n".join(chunks), document.page_count


def validate_current_preview() -> None:
    state = strict_json_load(STATE)
    matrix = strict_json_load(STATUS)
    manifest = strict_json_load(MANIFEST)
    if not isinstance(state, dict) or not isinstance(matrix, dict) or not isinstance(manifest, dict):
        raise AssertionError("MCL state, matrix and manifest must be JSON objects")
    if state["state"] != "PREVIEW" or state["publication_authority"] is not False:
        raise AssertionError("MCL v2.0 release state is not fail-closed PREVIEW")
    if manifest["status"] != "WORKING_PREVIEW" or manifest["publication_authority"] is not False:
        raise AssertionError("Working manifest claims publication authority")

    expected_inputs = {path.relative_to(ROOT).as_posix(): path for path in CONTROLLED_INPUTS}
    if set(manifest.get("controlled_inputs", {})) != set(expected_inputs):
        raise AssertionError("Working manifest controlled-input set is incomplete or unexpected")
    for relative_path, path in expected_inputs.items():
        if manifest["controlled_inputs"][relative_path] != {"bytes": path.stat().st_size, "sha256": sha256(path)}:
            raise AssertionError(f"Working manifest controlled-input mismatch: {relative_path}")

    expected_names = {
        "index.html",
        "guideline-working.docx",
        "guideline-working.pdf",
        "quickref-working.docx",
        "quickref-working.pdf",
        "algorithm-working.svg",
        "algorithm-working.excalidraw",
    }
    if set(manifest["artefacts"]) != expected_names:
        raise AssertionError("Working manifest artefact set is incomplete or unexpected")
    for name, record in manifest["artefacts"].items():
        path = PREVIEW / name
        if not path.is_file():
            raise AssertionError(f"Missing working artefact: {name}")
        if path.stat().st_size != record["bytes"] or sha256(path) != record["sha256"]:
            raise AssertionError(f"Working manifest mismatch: {name}")

    html_text = (PREVIEW / "index.html").read_text(encoding="utf-8")
    guideline_text = docx_text(PREVIEW / "guideline-working.docx")
    quickref_text = docx_text(PREVIEW / "quickref-working.docx")
    guideline_pdf_text, guideline_pages = pdf_text_and_pages(PREVIEW / "guideline-working.pdf")
    quickref_pdf_text, quickref_pages = pdf_text_and_pages(PREVIEW / "quickref-working.pdf")
    svg_text = (PREVIEW / "algorithm-working.svg").read_text(encoding="utf-8")
    excalidraw_text = (PREVIEW / "algorithm-working.excalidraw").read_text(encoding="utf-8")

    html_doc = BeautifulSoup(html_text, "html.parser")
    robots = html_doc.find_all("meta", attrs={"name": "robots"})
    if len(robots) != 1 or robots[0].get("content") != "noindex, nofollow, noarchive":
        raise AssertionError("HTML must contain exactly one controlled preview robots directive")
    evidence_ledger = strict_json_load(ROOT / "docs" / "mcl-v2" / "evidence-ledger.json")
    if not isinstance(evidence_ledger, dict) or len(html_doc.select(".evidence-reference")) != evidence_ledger.get("record_count"):
        raise AssertionError("HTML does not carry every controlled evidence reference")
    if html_doc.select_one("#release-control") is None or "Publication authority" not in html_text or "FALSE" not in html_text:
        raise AssertionError("HTML release-control record is absent or incomplete")
    warning = html_doc.select_one("aside.warning")
    if warning is None or warning.get("role") != "note":
        raise AssertionError("Static HTML preview warning must use a non-disruptive note role")

    if guideline_pages < 8:
        raise AssertionError(f"Guideline PDF is unexpectedly short: {guideline_pages} pages")
    if quickref_pages != 3:
        raise AssertionError(f"Quick-reference PDF must be exactly 3 pages, got {quickref_pages}")

    forbidden = ("AUTHORISED FOR PUBLICATION", "PUBLISHED EDUCATIONAL GUIDELINE", "PUBLISHED v2.0", "[VERIFY]")
    for label, text in (
        ("HTML", html_text),
        ("guideline DOCX", guideline_text),
        ("guideline PDF", guideline_pdf_text),
        ("quick-reference DOCX", quickref_text),
        ("quick-reference PDF", quickref_pdf_text),
        ("algorithm SVG", svg_text),
        ("algorithm Excalidraw", excalidraw_text),
    ):
        for phrase in forbidden:
            if phrase in text:
                raise AssertionError(f"{label} contains forbidden publication or unresolved wording: {phrase}")

    required_warning = "CONTROLLED WORKING PREVIEW"
    for label, text in (
        ("HTML", html_text),
        ("guideline DOCX", guideline_text),
        ("guideline PDF", guideline_pdf_text),
        ("quick-reference DOCX", quickref_text),
        ("quick-reference PDF", quickref_pdf_text),
        ("algorithm SVG", svg_text),
        ("algorithm Excalidraw", excalidraw_text),
    ):
        if required_warning not in text:
            raise AssertionError(f"{label} lacks the working-preview warning")

    source = BeautifulSoup(SOURCE.read_text(encoding="utf-8"), "html.parser")
    headings = [heading.get_text(" ", strip=True) for heading in source.select("section[data-clinical-unit] > h2")]
    if len(headings) != 12:
        raise AssertionError(f"Expected 12 canonical clinical headings, got {len(headings)}")
    for heading in headings:
        if heading not in html_text or heading not in guideline_text or heading not in guideline_pdf_text:
            raise AssertionError(f"Cross-format heading missing: {heading}")

    treatments = matrix["treatments"]
    if len(treatments) < 10 or len({item["id"] for item in treatments}) != len(treatments):
        raise AssertionError("Treatment-access matrix is incomplete or contains duplicate entries")
    for item in treatments:
        title = item["title"]
        if title not in html_text or title not in guideline_text or title not in guideline_pdf_text:
            raise AssertionError(f"Cross-format access entry missing: {title}")

    if not all(f"PART {number} OF 3" in quickref_text and f"PART {number} OF 3" in quickref_pdf_text for number in (1, 2, 3)):
        raise AssertionError("Quick reference does not preserve all three controlled parts")

    root = ET.fromstring(svg_text)
    if root.attrib.get("viewBox") != "0 0 1000 2050":
        raise AssertionError("Algorithm SVG viewBox changed unexpectedly")
    excalidraw = json.loads(excalidraw_text)
    if excalidraw.get("type") != "excalidraw" or len(excalidraw.get("elements", [])) < 40:
        raise AssertionError("Editable algorithm is malformed or incomplete")
    expected_semantics = {key: " ".join(value.split()) for key, value in canonical_semantic_texts().items()}
    svg_semantics = {
        element.attrib["data-semantic-id"]: " ".join(" ".join(element.itertext()).split())
        for element in root.iter()
        if "data-semantic-id" in element.attrib
    }
    editable_semantics = {
        element["customData"]["semanticId"]: " ".join(str(element["text"]).split())
        for element in excalidraw.get("elements", [])
        if element.get("type") == "text" and element.get("customData", {}).get("semanticId")
    }
    if svg_semantics != expected_semantics or editable_semantics != expected_semantics:
        raise AssertionError("SVG and Excalidraw are not exact renderings of the canonical semantic model")

    with tempfile.TemporaryDirectory(prefix="mcl-v2-repro-") as directory:
        regenerated = build_multiformat_preview(Path(directory))
        if regenerated["html"].read_bytes() != (PREVIEW / "index.html").read_bytes():
            raise AssertionError("HTML preview is not byte-reproducible")
        if regenerated["algorithm_svg"].read_bytes() != (PREVIEW / "algorithm-working.svg").read_bytes():
            raise AssertionError("Algorithm SVG is not byte-reproducible")
        if regenerated["algorithm_excalidraw"].read_bytes() != (PREVIEW / "algorithm-working.excalidraw").read_bytes():
            raise AssertionError("Algorithm Excalidraw is not byte-reproducible")
        if regenerated["guideline_docx"].read_bytes() != (PREVIEW / "guideline-working.docx").read_bytes():
            raise AssertionError("Guideline DOCX is not byte-reproducible")
        if regenerated["quickref_docx"].read_bytes() != (PREVIEW / "quickref-working.docx").read_bytes():
            raise AssertionError("Quick-reference DOCX is not byte-reproducible")
        regenerated_guideline_text, regenerated_guideline_pages = pdf_text_and_pages(regenerated["guideline_pdf"])
        regenerated_quickref_text, regenerated_quickref_pages = pdf_text_and_pages(regenerated["quickref_pdf"])
        if (regenerated_guideline_text, regenerated_guideline_pages) != (guideline_pdf_text, guideline_pages):
            raise AssertionError("Guideline PDF is not semantically reproducible")
        if (regenerated_quickref_text, regenerated_quickref_pages) != (quickref_pdf_text, quickref_pages):
            raise AssertionError("Quick-reference PDF is not semantically reproducible")

    print(
        "MCL v2.0 working-preview validation: PASS "
        f"artefacts=7 clinical_units=12 treatments={len(treatments)} guideline_pages={guideline_pages} quickref_pages={quickref_pages}"
    )


def main() -> int:
    validate_current_preview()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
