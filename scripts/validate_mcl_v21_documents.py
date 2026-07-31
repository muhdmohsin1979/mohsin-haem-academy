#!/usr/bin/env python3
"""Validate MCL v2.1 reviewed DOCX/PDF artefacts and semantic completeness."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any, cast

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PREVIEW = ROOT / "docs" / "mcl-v2.1" / "web-preview"
MANIFEST = PREVIEW / "manifest-reviewed-documents.json"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def docx_text(path: Path) -> str:
    document = Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for section in document.sections:
        chunks.extend(paragraph.text for paragraph in section.header.paragraphs)
        chunks.extend(paragraph.text for paragraph in section.footer.paragraphs)
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def pdf_text_pages(path: Path) -> tuple[str, int]:
    text: list[str] = []
    with fitz.open(path) as document:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            number = page_index + 1
            page_text = cast(str, page.get_text())
            if not page_text.strip():
                raise AssertionError(f"Blank PDF page: {path.name} page {number}")
            for block in cast(list[list[Any]], page.get_text("blocks")):
                x0, y0, x1, y1 = (float(value) for value in block[:4])
                if x0 < -1 or y0 < -1 or x1 > page.rect.width + 1 or y1 > page.rect.height + 1:
                    raise AssertionError(f"Out-of-bounds PDF text: {path.name} page {number}")
            text.append(page_text)
        return "\n".join(text), document.page_count


def validate() -> None:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    if manifest["publication_authority"] is not False:
        raise AssertionError("Reviewed documents claim publication authority")
    for name, record in manifest["artefacts"].items():
        path = PREVIEW / name
        if not path.is_file() or path.stat().st_size != record["bytes"] or sha256(path) != record["sha256"]:
            raise AssertionError(f"Reviewed document manifest mismatch: {name}")

    guideline_docx = docx_text(PREVIEW / "guideline-v2.1-reviewed.docx")
    quick_docx = docx_text(PREVIEW / "quickref-v2.1-reviewed.docx")
    guideline_pdf, guideline_pages = pdf_text_pages(PREVIEW / "guideline-v2.1-reviewed.pdf")
    quick_pdf, quick_pages = pdf_text_pages(PREVIEW / "quickref-v2.1-reviewed.pdf")

    for number in range(1, 22):
        marker = f"{number}. "
        if marker not in guideline_docx or marker not in guideline_pdf:
            raise AssertionError(f"Section {number} missing from guideline document family")
    for text in (guideline_docx, guideline_pdf):
        if text.count("Integrity:") != 53:
            raise AssertionError("Guideline document family does not preserve 53 evidence integrity records")
        normalised = " ".join(text.split())
        for phrase in (
            "CONTROLLED REVIEWED PREVIEW — NOT FOR CLINICAL USE",
            "Independent clinical review",
            "PASS — reviewer identity retained privately",
            "COMPLETE — verifier identity retained privately",
            "Publication authority",
            "FALSE — complete release family and exact production-hash ratification pending",
        ):
            if phrase not in normalised:
                raise AssertionError(f"Guideline document missing release wording: {phrase}")

    for text in (quick_docx, quick_pdf):
        normalised = " ".join(text.split())
        for phrase in (
            "MCL v2.1 quick reference",
            "One-page quick reference",
            "Trial referral",
            "Allogeneic HCT after CAR-T failure",
            "CONTROLLED REVIEWED PREVIEW",
        ):
            if phrase not in normalised:
                raise AssertionError(f"Quick-reference document missing: {phrase}")

    if guideline_pages < 20:
        raise AssertionError(f"Guideline PDF unexpectedly short: {guideline_pages}")
    if not 1 <= quick_pages <= 8:
        raise AssertionError(f"Quick-reference PDF page count is implausible: {quick_pages}")

    print(
        "MCL v2.1 reviewed documents: PASS "
        f"guideline_pages={guideline_pages} quickref_pages={quick_pages} evidence_records=53 publication_authority=false"
    )


if __name__ == "__main__":
    validate()
