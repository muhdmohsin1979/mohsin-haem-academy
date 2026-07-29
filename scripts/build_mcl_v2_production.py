#!/usr/bin/env python3
"""Build the MCL v2.0 production candidate from the frozen reviewed preview."""

from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as DocxDocument

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from scripts.build_mcl_v2_preview import artefact_record
from scripts.generate_mcl_v2_documents import canonicalise_docx
from scripts.generate_mcl_v2_release import apply_cll_navigation_pattern

PREVIEW = ROOT / "docs" / "mcl-v2" / "preview"
OUTPUT = ROOT / "guidelines" / "mcl"
REVIEWED_COMMIT = "763863cf3b89d207fce3bd29c4df076e0386ab92"
REVIEWED_TREE = "e4987ea5d1ee89bf0d9b307062299eb106d26af0"
REVIEWED_MANIFEST_SHA256 = "be160f955203e33b3a72e4b9829328358568bb8ba5337eb9922fb2f8bdff95fb"
PUBLICATION_DATE = "29 July 2026"

NAMES = {
    "index.html": "index.html",
    "guideline-working.docx": "guideline-v2.0.docx",
    "quickref-working.docx": "quickref-v2.0.docx",
    "algorithm-working.svg": "algorithm-v2.0.svg",
    "algorithm-working.excalidraw": "algorithm-v2.0.excalidraw",
}

REPLACEMENTS = {
    "CONTROLLED WORKING PREVIEW — NOT FOR CLINICAL USE": "AUTHORISED FOR PUBLICATION — 29 JULY 2026",
    "WORKING PREVIEW v2.0": "PUBLISHED v2.0",
    "non-public working preview": "published specialist guideline",
    "working preview": "published guideline",
    "working draft": "specialist guideline",
    "Non-public specialist clinical working draft": "Published specialist clinical guideline",
    "Clinical, pharmacy and publication review pending. No publication authority.": "Independent clinical review and pharmacy verification complete. Publication authorised.",
    "Clinical and pharmacy review pending. No publication authority. Do not use for treatment, prescribing, consent, referral or commissioning decisions.": "Independent clinical review and pharmacy verification complete. Use with the current SmPC, local SACT protocol and jurisdiction-specific commissioning criteria.",
    "This build is generated from controlled source material in PREVIEW state. Clinical review, pharmacy verification, publication review, owner approval, preview verification and publication authorisation remain pending.": "This published release was generated from the controlled source and approved after independent clinical review, pharmacy verification and exact-hash owner authorisation.",
    "Clinical and pharmacy review pending. This quick reference does not contain prescribing schedules and cannot replace the SmPC or local SACT protocol.": "This quick reference does not contain all prescribing schedules and cannot replace the current SmPC or local SACT protocol.",
    "Human pharmacy verification is mandatory before any publication candidate is frozen.": "Pharmacy verification was completed for the controlled release; use current SmPCs and local SACT protocols at the point of care.",
    "Clinical and pharmacy review pending — no publication authority. Recheck official national sources before treatment and before any release decision.": "Recheck current official national sources, SmPCs and local protocols before treatment because access and prescribing requirements may change.",
    "PREVIEW": "PUBLISHED",
    "PENDING": "COMPLETE",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def canonicalise_pdf(path: Path) -> None:
    with fitz.open(path) as document, tempfile.NamedTemporaryFile(
        prefix=f"{path.stem}-", suffix=".pdf", dir=path.parent, delete=False
    ) as handle:
        temporary = Path(handle.name)
        metadata = document.metadata or {}
        document.set_metadata({
            "title": metadata.get("title", ""),
            "author": metadata.get("author", ""),
            "subject": metadata.get("subject", ""),
            "keywords": metadata.get("keywords", ""),
            "creator": "Mohsin Haematology Academy controlled build",
            "producer": "Mohsin Haematology Academy controlled build",
            "creationDate": "D:20260729000000Z",
            "modDate": "D:20260729000000Z",
        })
        document.save(temporary, garbage=4, deflate=True, no_new_id=True)
    temporary.replace(path)


def convert_docx(path: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice soffice is required to generate production PDFs")
    with tempfile.TemporaryDirectory(prefix="mcl-v2-production-libreoffice-") as profile:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={(Path(profile) / 'profile').resolve().as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(path.parent),
                str(path),
            ],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
    output = path.with_suffix(".pdf")
    if result.returncode != 0 or not output.is_file():
        raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr or result.stdout}")
    canonicalise_pdf(output)
    return output


def frozen_pdf_for_docx(docx_path: Path, previous_artefacts: dict[str, object]) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    previous = previous_artefacts.get(pdf_path.name)
    if (
        pdf_path.is_file()
        and isinstance(previous, dict)
        and previous.get("source_docx_sha256") == sha256(docx_path)
        and previous.get("sha256") == sha256(pdf_path)
    ):
        return pdf_path
    return convert_docx(docx_path)


def assert_reviewed_preview() -> dict[str, object]:
    manifest_path = PREVIEW / "build-manifest-working.json"
    if sha256(manifest_path) != REVIEWED_MANIFEST_SHA256:
        raise AssertionError("Reviewed preview manifest hash changed")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    for name, record in manifest["artefacts"].items():
        path = PREVIEW / name
        if not path.is_file() or sha256(path) != record["sha256"]:
            raise AssertionError(f"Reviewed preview artefact changed: {name}")
    return manifest


def all_paragraphs(document: DocxDocument):
    for paragraph in document.paragraphs:
        yield paragraph
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def transform_docx(source: Path, target: Path) -> None:
    document = Document(str(source))
    for paragraph in all_paragraphs(document):
        text = paragraph.text
        updated = text
        for old, new in REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != text:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated
    properties = document.core_properties
    properties.title = properties.title.replace("working preview", "published guideline")
    properties.subject = "Published specialist clinical guideline"
    properties.keywords = properties.keywords.replace("working preview", "published guideline")
    properties.comments = "Independent clinical review and pharmacy verification complete. Publication authorised 29 July 2026."
    document.save(str(target))
    canonicalise_docx(target)


def transform_html(source: Path, target: Path) -> None:
    text = source.read_text(encoding="utf-8")
    replacements = {
        '<meta name="robots" content="noindex, nofollow, noarchive">': '<meta name="robots" content="index, follow">',
        "MCL v2.0 controlled working preview": "Mantle Cell Lymphoma guideline v2.0",
        '<span class="preview">CONTROLLED WORKING PREVIEW — NOT FOR CLINICAL USE</span>': '<span class="preview">PUBLISHED 29 JULY 2026</span>',
        '<strong id="preview-warning-heading">Non-public working build.</strong>\n    Clinical and pharmacy review are pending. This preview has no publication authority and must not be used for treatment, prescribing, consent, referral or commissioning decisions.': '<strong id="preview-warning-heading">Published specialist guideline.</strong>\n    Independent clinical review and pharmacy verification are complete. Use current SmPCs, local SACT protocols and jurisdiction-specific commissioning criteria.',
        'aria-label="MCL preview sections"': 'aria-label="MCL guideline sections"',
        "Working review artefacts": "Guideline downloads",
        "All files below are non-public working previews. They have no publication authority and remain subject to clinical, pharmacy and publication review.": "Controlled publication files for specialist healthcare professionals.",
        "guideline-working.pdf": "guideline-v2.0.pdf",
        "guideline-working.docx": "guideline-v2.0.docx",
        "quickref-working.pdf": "quickref-v2.0.pdf",
        "quickref-working.docx": "quickref-v2.0.docx",
        "algorithm-working.svg": "algorithm-v2.0.svg",
        "algorithm-working.excalidraw": "algorithm-v2.0.excalidraw",
        "Open the MCL v2.0 working treatment algorithm": "Open the MCL v2.0 treatment algorithm",
        "Evidence-first treatment pathway with explicit access labels. Working preview only.": "Evidence-first treatment pathway with explicit access labels.",
        "<tr><th scope=\"row\">Owner scope approval</th><td>FALSE</td></tr>": "<tr><th scope=\"row\">Owner scope approval</th><td>TRUE</td></tr>",
        "<tr><th scope=\"row\">Independent clinical review</th><td>PENDING</td></tr>": "<tr><th scope=\"row\">Independent clinical review</th><td>PASS</td></tr>",
        "<tr><th scope=\"row\">Pharmacy verification</th><td>PENDING</td></tr>": "<tr><th scope=\"row\">Pharmacy verification</th><td>COMPLETE — IDENTITY RETAINED PRIVATELY</td></tr>",
        "<tr><th scope=\"row\">Publication authority</th><td>FALSE</td></tr>": "<tr><th scope=\"row\">Publication authority</th><td>TRUE — 29 JULY 2026</td></tr>",
        "controlled working preview": "published guideline",
    }
    for old, new in replacements.items():
        if old not in text:
            raise AssertionError(f"Expected reviewed HTML release label is missing: {old[:70]}")
        text = text.replace(old, new)
    text = apply_cll_navigation_pattern(text, "MCL guideline sections")
    target.write_text(text, encoding="utf-8")


def transform_algorithm(source: Path, target: Path) -> None:
    text = source.read_text(encoding="utf-8")
    replacements = {
        '<tspan x="850" dy="0" font-weight="700">WORKING PREVIEW</tspan><tspan x="850" dy="22" font-weight="400">NOT FOR CLINICAL USE</tspan>': '<tspan x="850" dy="0" font-weight="700">PUBLISHED</tspan><tspan x="850" dy="22" font-weight="400">29 JULY 2026</tspan>',
        "WORKING PREVIEW\\nNOT FOR CLINICAL USE": "PUBLISHED 29 JULY 2026",
        "WORKING PREVIEW\nNOT FOR CLINICAL USE": "PUBLISHED 29 JULY 2026",
        "CONTROLLED WORKING PREVIEW \\u2014 do not use for treatment, prescribing, consent, referral or commissioning.": "PUBLISHED SPECIALIST GUIDELINE \\u2014 use current SmPCs, local protocols and commissioning criteria.",
        "CONTROLLED WORKING PREVIEW — do not use for treatment, prescribing, consent, referral or commissioning.": "PUBLISHED SPECIALIST GUIDELINE — use current SmPCs, local protocols and commissioning criteria.",
        "Clinical and pharmacy review pending \\u2014 no publication authority.": "Independent clinical review and release pharmacy verification complete.",
        "Clinical and pharmacy review pending — no publication authority.": "Independent clinical review and release pharmacy verification complete.",
        "Non-public working preview showing": "Published specialist algorithm showing",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    if "WORKING PREVIEW" in text or "NOT FOR CLINICAL USE" in text:
        raise AssertionError(f"Stale preview label remains in {source.name}")
    target.write_text(text, encoding="utf-8")


def build() -> dict[str, object]:
    assert_reviewed_preview()
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manifest_path = OUTPUT / "release-manifest-v2.0.json"
    previous_artefacts: dict[str, object] = {}
    if manifest_path.is_file():
        previous_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if isinstance(previous_manifest.get("artefacts"), dict):
            previous_artefacts = previous_manifest["artefacts"]
    transform_html(PREVIEW / "index.html", OUTPUT / "index.html")
    transform_docx(PREVIEW / "guideline-working.docx", OUTPUT / "guideline-v2.0.docx")
    transform_docx(PREVIEW / "quickref-working.docx", OUTPUT / "quickref-v2.0.docx")
    transform_algorithm(PREVIEW / "algorithm-working.svg", OUTPUT / "algorithm-v2.0.svg")
    transform_algorithm(PREVIEW / "algorithm-working.excalidraw", OUTPUT / "algorithm-v2.0.excalidraw")
    guideline_pdf = frozen_pdf_for_docx(OUTPUT / "guideline-v2.0.docx", previous_artefacts)
    quickref_pdf = frozen_pdf_for_docx(OUTPUT / "quickref-v2.0.docx", previous_artefacts)
    outputs = [
        OUTPUT / "index.html", OUTPUT / "guideline-v2.0.docx", guideline_pdf,
        OUTPUT / "quickref-v2.0.docx", quickref_pdf,
        OUTPUT / "algorithm-v2.0.svg", OUTPUT / "algorithm-v2.0.excalidraw",
    ]
    artefacts = {path.name: artefact_record(path) for path in outputs}
    artefacts[guideline_pdf.name]["source_docx_sha256"] = sha256(OUTPUT / "guideline-v2.0.docx")
    artefacts[quickref_pdf.name]["source_docx_sha256"] = sha256(OUTPUT / "quickref-v2.0.docx")
    manifest = {
        "schema_version": 1,
        "document_code": "MHA-MCL-2026-v2.0",
        "status": "AUTHORISED FOR PUBLICATION 29 JULY 2026",
        "publication_authority": True,
        "evidence_access_cut_off": "2026-07-28",
        "reviewed_preview": {
            "commit": REVIEWED_COMMIT,
            "tree": REVIEWED_TREE,
            "manifest_sha256": REVIEWED_MANIFEST_SHA256,
        },
        "clinical_change_from_reviewed_preview": "NONE; release-control presentation only",
        "artefacts": artefacts,
    }
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return manifest


if __name__ == "__main__":
    result = build()
    print(json.dumps(result, indent=2))
