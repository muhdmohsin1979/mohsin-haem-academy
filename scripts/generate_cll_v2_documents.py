#!/usr/bin/env python3
"""Generate controlled CLL v2.0 DOCX artefacts from the canonical HTML.

Dependencies: python-docx, beautifulsoup4. Convert DOCX to PDF with LibreOffice.
Release status is controlled by the manifest, owner authorisation and release record.
"""

from __future__ import annotations

import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
CLL = ROOT / "guidelines" / "cll"
HTML = CLL / "index.html"
GUIDELINE_DOCX = CLL / "guideline.docx"
QUICKREF_DOCX = CLL / "quickref.docx"

NAVY = "1B2A4A"
RED = "C41E3A"
LIGHT_BLUE = "EAF0F7"
LIGHT_RED = "FBECEF"
LIGHT_GREY = "F3F4F6"
MID_GREY = "6B7280"
BORDER = "D7DCE2"
DOC_CODE = "MHA-CLL-2026-v2.0"
CUT_OFF = "26 July 2026"

SUPER = str.maketrans("0123456789-+", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻⁺")


def node_text(node: Tag) -> str:
    parts: list[str] = []

    def walk(item: Tag | NavigableString) -> None:
        if isinstance(item, NavigableString):
            parts.append(str(item))
            return
        if item.name == "br":
            parts.append("\n")
            return
        if item.name == "sup":
            parts.append(item.get_text("", strip=True).translate(SUPER))
            return
        for child in item.children:
            walk(child)
        if item.name in {"p", "li", "div", "td", "th"}:
            parts.append(" ")

    walk(node)
    text = "".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_document(doc: Document, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, colour in (("Title", 25, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 13, RED), ("Heading 3", 10.5, NAVY)):
        style = doc.styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    if "MHA Callout" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("MHA Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Georgia"
        style.font.size = Pt(8.8)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Cm(0.3)
        style.paragraph_format.right_indent = Cm(0.3)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(5.9)
    table.cell(0, 0).text = "MOHSIN HAEMATOLOGY ACADEMY"
    table.cell(0, 1).text = "PUBLISHED v2.0"
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, NAVY if idx == 0 else RED)
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = f"{DOC_CODE}  |  {subtitle}  |  Evidence cut-off {CUT_OFF}  |  Published 27 July 2026"
    p.style = doc.styles["Normal"]
    p.runs[0].font.size = Pt(7.5)
    p.runs[0].font.color.rgb = RGBColor.from_string(MID_GREY)
    add_page_number(footer.add_paragraph())

    props = doc.core_properties
    props.title = f"CLL v2.0 — {subtitle}"
    props.subject = "Published educational clinical decision support"
    props.author = "Dr Muhammad Mohsin, Consultant Haematologist"
    props.keywords = "CLL, relapsed refractory, pirtobrutinib, NICE TA1173"
    props.comments = "Published 27 July 2026 after independent review, pharmacy verification and clinical-owner authorisation."
    controlled_time = datetime(2026, 7, 26, 12, 0, tzinfo=timezone.utc)
    props.created = controlled_time
    props.modified = controlled_time


def add_release_banner(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_RED)
    set_cell_border(cell, RED, 8)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PUBLISHED EDUCATIONAL GUIDELINE — v2.0 — 27 JULY 2026")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(RED)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(subtitle)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(RED)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(f"{DOC_CODE}\nEvidence and access cut-off: {CUT_OFF}\nClinical owner: Dr Muhammad Mohsin, Consultant Haematologist")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_release_banner(doc)
    p = doc.add_paragraph(style="MHA Callout")
    p.add_run("Clinical decision-support only. ").bold = True
    p.add_run("Not for direct patient use. Use current SmPCs, NICE guidance, NHS commissioning rules, local policy and patient-specific clinical judgement. England-specific NICE/NHS access must not be described as a single UK-wide entitlement.")


def add_text_paragraph(doc: Document, text: str, style: str | None = None, bold_lead: str | None = None) -> None:
    if not text:
        return
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def render_list(doc: Document, node: Tag, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        text = node_text(li)
        if text:
            p = doc.add_paragraph(text, style=style)
            p.paragraph_format.space_after = Pt(2)


def render_table(doc: Document, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    if max_cols == 0:
        return
    table = doc.add_table(rows=0, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row_node in enumerate(rows):
        cells = row_node.find_all(["th", "td"], recursive=False)
        row = table.add_row()
        keep_row_together(row)
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            text = node_text(cells[c_idx]) if c_idx < len(cells) else ""
            cell.text = text
            if r_idx == 0 or (c_idx < len(cells) and cells[c_idx].name == "th"):
                set_cell_shading(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(8)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, LIGHT_GREY)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(7.8)
        if r_idx == 0:
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph()


def render_callout(doc: Document, node: Tag, fill: str = LIGHT_BLUE) -> None:
    heading = node.select_one(".info-box-head, .uk-box-title, .col-card-head, .tx-card-title")
    body = node.select_one(".info-box-body, .col-card-body, .tx-card-body")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, RED if fill == LIGHT_RED else NAVY, 6)
    if heading:
        p = cell.paragraphs[0]
        p.text = node_text(heading)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(RED if fill == LIGHT_RED else NAVY)
            run.font.size = Pt(9.3)
    target = body or node
    block_children = [
        child for child in target.children
        if isinstance(child, Tag) and child.name in {"p", "ul", "ol", "table"}
    ]
    if not block_children:
        text = node_text(target)
        if text:
            p = cell.add_paragraph(text)
            p.style = doc.styles["MHA Callout"]
    for child in target.children:
        if not isinstance(child, Tag):
            continue
        if child.name == "p":
            p = cell.add_paragraph(node_text(child))
            p.style = doc.styles["MHA Callout"]
        elif child.name in {"ul", "ol"}:
            for li in child.find_all("li", recursive=False):
                p = cell.add_paragraph(node_text(li), style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
        elif child.name == "table":
            p = cell.add_paragraph(node_text(child))
            p.style = doc.styles["MHA Callout"]
    doc.add_paragraph()


def render_container(doc: Document, container: Tag) -> None:
    for child in container.children:
        if not isinstance(child, Tag):
            continue
        classes = set(child.get("class", []))
        if child.name in {"svg", "object", "figure", "script", "style"}:
            continue
        if "section-header" in classes:
            continue
        if child.name == "h2":
            p = doc.add_paragraph(node_text(child), style="Heading 1")
            keep_with_next(p)
        elif child.name == "h3":
            p = doc.add_paragraph(node_text(child), style="Heading 2")
            keep_with_next(p)
        elif child.name in {"h4", "h5"}:
            p = doc.add_paragraph(node_text(child), style="Heading 3")
            keep_with_next(p)
        elif child.name == "p":
            add_text_paragraph(doc, node_text(child))
        elif child.name == "ul":
            render_list(doc, child)
        elif child.name == "ol":
            render_list(doc, child, ordered=True)
        elif child.name == "table":
            render_table(doc, child)
        elif "tx-card" in classes:
            render_callout(doc, child, LIGHT_BLUE)
        elif classes.intersection({"info-box", "uk-box", "keypoint", "edu-banner"}):
            fill = LIGHT_RED if classes.intersection({"box-danger", "box-warn"}) else LIGHT_BLUE
            render_callout(doc, child, fill)
        elif "vignette" in classes:
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, LIGHT_GREY)
            set_cell_border(cell)
            cell.text = node_text(child)
            cell.paragraphs[0].runs[0].italic = True
            doc.add_paragraph()
        elif "two-col" in classes:
            for card in child.find_all(class_=re.compile(r"col-card$"), recursive=False):
                render_callout(doc, card, LIGHT_BLUE)
        elif child.name in {"div", "details"}:
            if "hero" in " ".join(classes) or "tool-meta" in classes:
                continue
            render_container(doc, child)


def build_guideline(soup: BeautifulSoup) -> None:
    doc = Document()
    configure_document(doc, "Full guideline")
    add_title(doc, "Chronic Lymphocytic Leukaemia", "Unified clinical guideline — v2.0 relapsed/refractory update")

    doc.add_heading("Release status", level=1)
    add_text_paragraph(doc, "Published 27 July 2026 after independent review, pharmacy verification and clinical-owner approval of the manifest-bound release candidate. This is educational clinical decision support, not a prescribing protocol.")

    doc.add_heading("Quick decision summary", level=1)
    quick = soup.select_one(".gl-main > div[style*='border-left:4px']")
    if quick:
        ul = quick.find("ul")
        if ul:
            render_list(doc, ul)

    main = soup.select_one("main.gl-main")
    if main is None:
        raise RuntimeError("Canonical main content not found")
    for section in main.find_all("section", class_="gl-section", recursive=False):
        section_id = section.get("id", "")
        if section_id in {"cite", "governance"}:
            continue
        if not section_id and "Versioning & Governance" in node_text(section):
            continue
        heading = section.select_one(".section-title")
        if heading:
            p = doc.add_paragraph(node_text(heading), style="Heading 1")
            keep_with_next(p)
        render_container(doc, section)

    doc.add_heading("Version history", level=1)
    history = doc.add_table(rows=1, cols=4)
    for index, label in enumerate(["Version", "Date", "Author", "Change"]):
        history.rows[0].cells[index].text = label
        set_cell_shading(history.rows[0].cells[index], NAVY)
        set_cell_border(history.rows[0].cells[index])
        for run in history.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for values in [
        ["v1.0", "Apr 2025", "Dr M Mohsin", "Initial controlled publication."],
        ["v1.9", "Apr 2026", "Dr M Mohsin", "Published baseline before the substantive R/R revision."],
        ["v2.0", "27 Jul 2026", "Dr M Mohsin", "Published TA1173 and evidence-led R/R sequencing update."],
    ]:
        row = history.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            row.cells[index].text = value
            set_cell_border(row.cells[index])

    doc.add_heading("Release-control record", level=1)
    table = doc.add_table(rows=0, cols=2)
    records = [
        ("Document code", DOC_CODE),
        ("Evidence/access cut-off", CUT_OFF),
        ("Scope approval", "Approved by clinical owner, 26 July 2026"),
        ("Pharmacy verification", "COMPLETE — verifier identity retained privately"),
        ("Independent clinical/content review", "PASS"),
        ("Exact-artefact owner approval", "COMPLETE"),
        ("Publication", "AUTHORISED 27 JULY 2026"),
    ]
    for key, value in records:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GREY)
        for cell in row.cells:
            set_cell_border(cell)
    doc.save(GUIDELINE_DOCX)


def add_quickref_table(doc: Document, heading: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph(heading, style="Heading 2")
    keep_with_next(p)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, NAVY)
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(7.8)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.text = value
            set_cell_border(cell)
            if r_idx % 2:
                set_cell_shading(cell, LIGHT_GREY)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7.3)
    doc.add_paragraph()


def build_quickref() -> None:
    doc = Document()
    configure_document(doc, "Quick reference")
    add_title(doc, "CLL quick reference", "v2.0 relapsed/refractory update")

    doc.add_heading("Before any treatment line", level=1)
    for item in [
        "Treat only when iwCLL active-disease criteria are met.",
        "Review prior regimen, class, duration, depth of response, reason for stopping and treatment-free interval.",
        "Repeat del(17p) FISH and TP53 sequencing before the next treatment line.",
        "Assess rapidly progressive or asymmetric disease, disproportionate LDH rise or systemic deterioration for Richter transformation; use PET-CT-directed biopsy.",
        "Document fitness, cardiac and bleeding risk, renal function, TLS risk, infection burden, interactions, patient preference and access route.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_quickref_table(doc, "First-line orientation", ["State", "Route"], [
        ["del(17p) and/or TP53 mutation", "Targeted therapy only. BSH 2025 favours zanubrutinib or acalabrutinib; TA931 and TA689 apply only within their specified untreated populations. Use current SmPC and Blueteq criteria."],
        ["No TP53 abnormality", "Continuous zanubrutinib/acalabrutinib or fixed-duration venetoclax–obinutuzumab. Decide with the patient and apply the exact treatment-specific NICE criteria."],
    ])

    doc.add_page_break()
    doc.add_heading("Relapsed/refractory treatment state", level=1)
    add_quickref_table(doc, "Definitions", ["State", "Definition"], [
        ["Intolerant", "Stopped for toxicity without progression. A better-tolerated agent within class may remain appropriate."],
        ["Exposed, not refractory", "Completed or stopped treatment without progression; sensitivity may remain."],
        ["Refractory", "Progression during active treatment or no meaningful response."],
        ["Double exposed", "Both a covalent BTKi and venetoclax have been used."],
        ["Double refractory", "Progression during both classes. Distinct high-risk state; not synonymous with double exposure."],
    ])
    add_quickref_table(doc, "Sequencing", ["Prior-treatment state", "Decision route", "Do not"], [
        ["Covalent-BTKi intolerance", "Alternative better-tolerated covalent BTKi or different class according to comorbidity, preference and NICE eligibility.", "Do not describe an intolerance switch as efficacy after resistance."],
        ["Covalent-BTKi progression; venetoclax naive", "Venetoclax-based NICE pathway when eligible; pirtobrutinib under TA1173 when its restriction is met.", "Do not continue covalent BTKi monotherapy as definitive resistant-disease treatment."],
        ["Relapse after fixed-duration venetoclax response", "BTKi or selected venetoclax retreatment after reviewing prior response and treatment-free interval.", "Do not extrapolate retreatment evidence to primary or on-treatment resistance."],
        ["Venetoclax progression; BTKi naive", "NICE-funded covalent BTKi when eligible.", "Do not claim a universal randomised sequencing hierarchy."],
        ["Double exposed, not double refractory", "Individualise by retained sensitivity, prior response and access.", "Do not automatically label double refractory."],
        ["Double refractory", "Early tertiary/trial referral; pirtobrutinib if TA1173 criteria met; highly selected cellular therapy/alloHCT discussion.", "Do not imply routine NHS England CLL CAR-T commissioning or established devolved-nation access, or a comparative hierarchy."],
    ])

    doc.add_page_break()
    doc.add_heading("NICE recommendations and NHS England access at 26 July 2026", level=1)
    add_quickref_table(doc, "Access and schedule", ["Option", "NICE position", "Schedule / control"], [
        ["Venetoclax–rituximab", "TA561 after ≥1 previous therapy.", "Five-week venetoclax ramp-up to 400 mg OD before rituximab; rituximab 375 mg/m² C1D1 then 500 mg/m² D1 C2–6; venetoclax for 24 months from C1D1 rituximab."],
        ["Venetoclax monotherapy", "TA796: (a) del(17p)/TP53-mutated CLL when a B-cell receptor pathway inhibitor is unsuitable, or after progression on one; or (b) CLL without del(17p)/TP53 mutation after progression following both chemoimmunotherapy and a B-cell receptor pathway inhibitor.", "Five-week ramp-up to 400 mg OD; continue until progression or unacceptable toxicity."],
        ["Acalabrutinib", "TA689 for previously treated CLL and specified untreated populations.", "100 mg BD continuously; current SmPC and Blueteq criteria."],
        ["Zanubrutinib", "TA931 for R/R CLL and specified untreated populations.", "160 mg BD or 320 mg OD continuously; current SmPC and Blueteq criteria."],
        ["Pirtobrutinib", "TA1173: adults with R/R CLL who have had a BTKi, only if covalent-BTKi retreatment, including after a fixed-duration regimen, is not clinically appropriate.", "200 mg OD continuously. At 26 July 2026 NHS England access remained through interim Blueteq funding; routine-budget implementation was due by 29 September 2026. The broader MHRA licence does not widen NHS funding."],
    ])

    p = doc.add_paragraph(style="MHA Callout")
    p.add_run("Venetoclax safety: ").bold = True
    p.add_run("Apply current SmPC TLS stratification, prophylaxis, monitoring, renal-function assessment and interaction management. This card does not replace the prescribing information or local protocol.")

    doc.add_heading("Emerging or specialist options", level=2)
    for item in [
        "Pirtobrutinib–venetoclax–rituximab: phase III BRUIN CLL-322 evidence; NICE ID6566/GID-TA11748 awaiting development at the cut-off. Not routine NHS treatment.",
        "Lisocabtagene maraleucel: single-arm CLL evidence; no demonstrated CLL marketing authorisation, NICE recommendation or routine NHS England commissioning at the cut-off. Devolved-nation access was not established by this audit. NICE ID6174/GID-TA11001 discontinued February 2026.",
        "CD20×CD3 bispecific antibodies: no qualifying peer-reviewed prospective therapeutic trial in untransformed CLL identified through the cut-off. Do not extrapolate Richter or other-lymphoma evidence.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Release control", level=2)
    add_text_paragraph(doc, "Published 27 July 2026. Independent review, pharmacy verification and clinical-owner authorisation are complete. Educational clinical decision support only; verify current NICE, NHS, SmPC and local policy before prescribing.")
    doc.save(QUICKREF_DOCX)


def main() -> int:
    if not HTML.exists():
        print(f"Missing canonical HTML: {HTML}", file=sys.stderr)
        return 1
    soup = BeautifulSoup(HTML.read_text(encoding="utf-8"), "html.parser")
    build_guideline(soup)
    build_quickref()
    print(GUIDELINE_DOCX)
    print(QUICKREF_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
