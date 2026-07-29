from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
PREVIEW_STATE = ROOT / "tests" / "fixtures" / "mcl-release-state-preview.json"

from scripts.generate_mcl_v2_documents import build_documents


def document_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


class MCLV2DocumentGenerationTests(unittest.TestCase):
    def test_docx_bytes_and_package_metadata_are_canonical(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            first = build_documents(Path(directory) / "first", PREVIEW_STATE)
            second = build_documents(Path(directory) / "second", PREVIEW_STATE)
            for key in ("guideline", "quickref"):
                self.assertEqual(first[key].read_bytes(), second[key].read_bytes(), key)
                with zipfile.ZipFile(first[key]) as package:
                    names = package.namelist()
                    self.assertEqual(names, sorted(names))
                    self.assertFalse(any(name.startswith("customXml/") for name in names))
                    self.assertNotIn("docProps/thumbnail.jpeg", names)
                    self.assertEqual({item.date_time for item in package.infolist()}, {(2026, 7, 28, 7, 0, 0)})
                    app = package.read("docProps/app.xml").decode("utf-8")
                    self.assertIn("Mohsin Haematology Academy controlled generator", app)
                    self.assertNotIn("Microsoft Macintosh Word", app)
                    self.assertNotIn("Normal.dotm", app)
                    self.assertNotIn("<Pages>1</Pages>", app)
                    self.assertNotIn("<Words>0</Words>", app)

    def test_builds_non_public_guideline_and_three_part_quick_reference(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            outputs = build_documents(Path(directory), PREVIEW_STATE)

            guideline = outputs["guideline"]
            quickref = outputs["quickref"]
            self.assertTrue(guideline.is_file())
            self.assertTrue(quickref.is_file())

            guideline_text = document_text(guideline)
            quickref_text = document_text(quickref)

            self.assertIn("CONTROLLED WORKING PREVIEW — NOT FOR CLINICAL USE", guideline_text)
            self.assertIn("1. Scope and use", guideline_text)
            self.assertIn("13. Working regulatory and access matrix", guideline_text)
            self.assertIn("Evidence references", guideline_text)
            self.assertIn("VR-CAP / bortezomib combination", guideline_text)
            self.assertNotIn("AUTHORISED FOR PUBLICATION", guideline_text)
            self.assertNotIn("PUBLISHED EDUCATIONAL GUIDELINE", guideline_text)

            self.assertIn("PART 1 OF 3 — FIRST-LINE ORIENTATION", quickref_text)
            self.assertIn("PART 2 OF 3 — RELAPSED OR REFRACTORY PATHWAY", quickref_text)
            self.assertIn("PART 3 OF 3 — ACCESS AND SAFETY CONTROLS", quickref_text)
            self.assertIn("Clinical and pharmacy review pending", quickref_text)
            self.assertNotIn("AUTHORISED FOR PUBLICATION", quickref_text)

            quickref_doc = Document(quickref)
            self.assertGreaterEqual(quickref_doc.styles["Normal"].font.size.pt, 10)
            table_font_sizes = [
                run.font.size.pt
                for table in quickref_doc.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
                for run in paragraph.runs
                if run.text.strip() and run.font.size is not None
            ]
            self.assertTrue(table_font_sizes)
            self.assertGreaterEqual(min(table_font_sizes), 9)


if __name__ == "__main__":
    unittest.main()
